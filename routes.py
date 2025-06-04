import os
from dotenv import load_dotenv
from azure.core.credentials import AzureKeyCredential
from azure.ai.documentintelligence import DocumentIntelligenceClient
# from azure.ai.documentintelligence.models import AnalyzeResult
from flask import Flask, request, jsonify, send_from_directory, session
import openai
from PyPDF2 import PdfReader
import io
from azure.storage.blob import BlobServiceClient, BlobClient
import requests
import re
import base64
from urllib.parse import unquote
from docx import Document  # Add this import at the top
from docx2pdf import convert
import pythoncom

# Load environment variables
load_dotenv()

# Azure Configuration
AZURE_OAI_ENDPOINT = os.getenv("AZURE_OAI_ENDPOINT")
AZURE_OAI_KEY = os.getenv("AZURE_OAI_KEY")
AZURE_OAI_DEPLOYMENT = os.getenv("AZURE_OAI_DEPLOYMENT")
AZURE_DOC_INTELLIGENCE_ENDPOINT = os.getenv("AZURE_DOC_INTELLIGENCE_ENDPOINT")
AZURE_DOC_INTELLIGENCE_KEY = os.getenv("AZURE_DOC_INTELLIGENCE_KEY")

# Flask Configuration
app = Flask(__name__)
app.secret_key = os.getenv('FLASK_SECRET_KEY', 'your-secret-key-here')
app.config["UPLOAD_FOLDER"] = "uploads"
os.makedirs(app.config["UPLOAD_FOLDER"], exist_ok=True)

def analyze_layout_from_file(filepath):
    document_intelligence_client = DocumentIntelligenceClient(
        endpoint=AZURE_DOC_INTELLIGENCE_ENDPOINT,
        credential=AzureKeyCredential(AZURE_DOC_INTELLIGENCE_KEY)
    )
    
    # Read PDF and split into smaller chunks
    pdf_reader = PdfReader(filepath)
    total_pages = len(pdf_reader.pages)
    print(f"Processing {total_pages} pages...")
    
    all_results = []
    # Reduce chunk size to handle API limits
    chunk_size = 2  # Process 2 pages at a time
    
    for start_page in range(0, total_pages, chunk_size):
        end_page = min(start_page + chunk_size, total_pages)
        print(f"Processing pages {start_page + 1} to {end_page}...")
        
        # Create PDF for current chunk
        from PyPDF2 import PdfWriter
        pdf_writer = PdfWriter()
        for page_num in range(start_page, end_page):
            pdf_writer.add_page(pdf_reader.pages[page_num])
        
        # Save chunk to memory with proper cleanup
        pdf_bytes = io.BytesIO()
        pdf_writer.write(pdf_bytes)
        pdf_bytes.seek(0)
        chunk_content = pdf_bytes.read()
        pdf_bytes.close()
        
        try:
            # Analyze current chunk with explicit options
            poller = document_intelligence_client.begin_analyze_document(
                "prebuilt-layout",
                body=chunk_content or pdf_bytes,
                content_type="application/pdf"
            )
            result = poller.result()
            
            if result and hasattr(result, 'pages') and len(result.pages) > 0:
                all_results.append(result)
                print(f"Successfully processed pages {start_page + 1} to {end_page}")
            else:
                print(f"Warning: No content extracted from pages {start_page + 1} to {end_page}")
            
        except Exception as e:
            print(f"Error processing pages {start_page + 1} to {end_page}: {str(e)}")
            continue
    
    # Verify results
    if not all_results:
        raise Exception("Document processing failed: No pages were successfully processed")
    
    print(f"Successfully processed {len(all_results)} chunks")
    
    # Merge results into single result object
    final_result = all_results[0]
    for result in all_results[1:]:
        if hasattr(result, 'pages'):
            final_result.pages.extend(result.pages)
        if hasattr(result, 'tables') and result.tables:
            if not hasattr(final_result, 'tables'):
                final_result.tables = []
            final_result.tables.extend(result.tables)
    
    return final_result

def extract_text(result):
    all_text = []
    
    # Get total page count
    total_pages = len(result.pages)
    print(f"Extracting text from {total_pages} pages...")
    
    for page in result.pages:
        # Add clear page separator
        all_text.append(f"\n{'='*20} Page {page.page_number} of {total_pages} {'='*20}\n")
        
        # Extract all text content from the page
        lines = [line.content for line in page.lines]
        page_text = '\n'.join(lines)
        all_text.append(page_text)
        
        # Add tables if present
        if result.tables:
            tables_on_page = [
                table for table in result.tables 
                if any(region.page_number == page.page_number for region in table.bounding_regions)
            ]
            if tables_on_page:
                all_text.append("\n--- Tables on this page ---")
                for table in tables_on_page:
                    table_content = []
                    for cell in table.cells:
                        if cell.content.strip():
                            table_content.append(f"{cell.content}")
                    all_text.append(' | '.join(table_content))
    
    return '\n'.join(all_text)

def layout_to_text(result):
    output = []
    
    # Document properties
    if result.styles:
        output.append("Document Properties:")
        output.append("- Handwritten content: " + 
                     ("Yes" if any(style.is_handwritten for style in result.styles) else "No"))
    
    # Page analysis
    for page in result.pages:
        output.append(f"\n=== Page {page.page_number} Analysis ===")
        output.append(f"Dimensions: {page.width}x{page.height} {page.unit}")
        
        # Tables
        if result.tables:
            output.append("\nTables Found:")
            for table_idx, table in enumerate(result.tables):
                if table.bounding_regions and any(region.page_number == page.page_number 
                                                for region in table.bounding_regions):
                    output.append(f"Table {table_idx+1}: {table.row_count} rows x {table.column_count} columns")
                    table_content = []
                    for cell in table.cells:
                        if cell.bounding_regions and any(region.page_number == page.page_number 
                                                       for region in cell.bounding_regions):
                            table_content.append(f"Cell [{cell.row_index},{cell.column_index}]: {cell.content}")
                    output.extend(table_content)
    
    return "\n".join(output)

def summarize_text(text):
    client = openai.AzureOpenAI(
        azure_endpoint=AZURE_OAI_ENDPOINT,
        api_key=AZURE_OAI_KEY,
        api_version="2024-12-01-preview"
    )
    response = client.chat.completions.create(
        model=AZURE_OAI_DEPLOYMENT,
        messages=[
            {"role": "system", "content": (
                (
                    "You are an AI assistant. You are an RFP analysis assistant. When you answer, use HTML for formatting: "
                    "- Use <b> for bold, <ul>/<li> for lists, and <p> for paragraphs. When summarizing documents, start with 'Here is the summary of the RFP document:' and provide a clear, structured summary focusing on key requirements, deadlines, and specifications. "
                    "Pain Points. "
                    "Focus on: Technology Requirements, User Numbers, Deployment Scale, Current Systems, Pain Points, Required Integrations, Platforms, Compliance Requirements. "
                    "You also need to cite the section of the document where you found the information. "
                    "Technology Requirements: User Numbers: Deployment Scale: Current Systems: Pain Points: Required Integrations: Platforms: Compliance Requirements: Actionable Tasks for Proposal Team. "
                    "You also need to cite the section of the document where you found the information. "
                    "Purpose Scope Includes: Platform Required: Bid Deadline: Non-compulsory Briefing Target Contract Start Date: Contract Duration: Evaluation Process. "
                    "Your purpose is to help the sales team understand the document quickly."
                )
            )},
            {"role": "assistant", "content": "I have analyzed this RFP/RFQ document."},
            {"role": "user", "content": text}
        ],
        max_completion_tokens=4000  
    )
    return response.choices[0].message.content

def markdown_to_html(text):
    # Remove asterisks used for bullet points
    text = re.sub(r'^\s*[\*\-]\s+', '&bull; ', text, flags=re.MULTILINE)
    # Bold: **text** or __text__ to <b>text</b>
    text = re.sub(r'\*\*(.*?)\*\*', r'<b>\1</b>', text)
    text = re.sub(r'__(.*?)__', r'<b>\1</b>', text)
    # New paragraphs for double newlines
    text = re.sub(r'\n\s*\n', '</p><p>', text)
    # Single newline to <br>
    text = re.sub(r'(?<!</p>)\n', '<br>', text)
    # Wrap in <p>
    text = f'<p>{text}</p>'
    # Highlight dates (simple pattern)
    text = re.sub(r'(\d{1,2} [A-Za-z]+ \d{4})', r'<span class="highlight-date">\1</span>', text)
    # Highlight emails
    text = re.sub(r'([a-zA-Z0-9_.+-]+@[a-zA-Z0-9-]+\.[a-zA-Z0-9-.]+)', r'<span class="highlight-email">\1</span>', text)
    return text

@app.route("/")
def home():
    return send_from_directory('.', 'new 1.html')

@app.route("/upload", methods=["POST"])
def upload_file():
    if "document" not in request.files:
        return jsonify({"error": "No file uploaded"}), 400
    
    file = request.files["document"]
    if not file.filename:
        return jsonify({"error": "No file selected"}), 400

    ext = os.path.splitext(file.filename)[1].lower()
    file_path = os.path.join(app.config["UPLOAD_FOLDER"], file.filename)
    file.save(file_path)

    try:
        if ext == ".pdf":
            result = analyze_layout_from_file(file_path)
            extracted_text = extract_text(result)
            summary = summarize_text(extracted_text) if extracted_text else "No text extracted from document."
            missing_fields = detect_missing_fields(extracted_text)
            layout = layout_to_text(result)
        elif ext == ".docx":
            try:
                # Convert DOCX to PDF
                pdf_path = convert_word_to_pdf(file_path)
                # Process as PDF
                result = analyze_layout_from_file(pdf_path)
                extracted_text = extract_text(result)
                summary = summarize_text(extracted_text) if extracted_text else "No text extracted from document."
                missing_fields = detect_missing_fields(extracted_text)
                layout = layout_to_text(result)
                # Clean up generated PDF
                os.remove(pdf_path)
            except Exception as e:
                return jsonify({"error": f"Failed to process DOCX: {str(e)}"}), 400
        else:
            return jsonify({"error": "Unsupported file type. Please upload a PDF or DOCX."}), 400

        session['document_text'] = extracted_text

        return jsonify({
            "document_text": extracted_text,
            "summary": summary,
            "layout": layout,
            "missing_fields": missing_fields
        })
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route("/chat", methods=["POST"])
def chatbot_route():
    data = request.get_json()
    user_input = data['message']
    document_text = data.get('document_text', '')
    ai_search_passages = data.get('ai_search_passages', [])

    # Combine both sources for context
    context = ""
    if document_text:
        context += "[UPLOADED DOCUMENT]\n" + document_text[:2000]  # limit for token safety
    if ai_search_passages:
        context += "\n[AI SEARCH PASSAGES]\n" + "\n\n".join(ai_search_passages)

    # Define system_prompt BEFORE using it
    system_prompt = ("""You are an AI assistant helping Think Tank Sales Team analyze RFP/RFQ documents. Always provide specific, relevant information and cite the section of the document where you found the information. You are an RFP analysis assistant. When you answer, use HTML for formatting: - Use <b> for bold, <ul>/<li> for lists, and <p> for paragraphs. - Do not use asterisks or markdown. - Make your answers readable and well-structured. Focus on: 1. Technology requirements and specifications 2. User numbers and deployment scale 3. Current systems and pain points 4. Required integrations and platforms 5. Matching with past proposals and BOMs 6. Identifying compliance requirements 7. Suggesting actionable tasks for the proposal team you also need to cite the section of the document where you found the information. If suggesting BOM items, explain why they match the requirements."""
        
    )

    messages = [
        {"role": "system", "content": system_prompt},
        {"role": "assistant", "content": f"Context:\n{context}"},
        {"role": "user", "content": user_input}
    ]

    client = openai.AzureOpenAI(
        azure_endpoint=AZURE_OAI_ENDPOINT,
        api_key=AZURE_OAI_KEY,
        api_version="2024-12-01-preview"
    )
    response = client.chat.completions.create(
        model=AZURE_OAI_DEPLOYMENT,
        messages=messages,
        max_completion_tokens=4000
    )

    bot_response = response.choices[0].message.content
    bot_response_html = markdown_to_html(bot_response)
    return jsonify({"response": bot_response_html})

@app.route("/search/list", methods=["POST"])
def list_rfp_files():
    data = request.get_json()
    search_query = data.get("search", "*")
    search_endpoint = os.getenv("AZURE_SEARCH_ENDPOINT")
    search_key = os.getenv("AZURE_SEARCH_KEY")
    search_index = os.getenv("AZURE_SEARCH_INDEX")
    blob_container = os.getenv("AZURE_BLOB_CONTAINER")

    url = f"{search_endpoint}/indexes/{search_index}/docs/search?api-version=2023-07-01-Preview"
    headers = {
        "Content-Type": "application/json",
        "api-key": search_key
    }
    payload = {
        "search": "*",  # Always fetch all files
        "top": 1000
    }

    try:
        resp = requests.post(url, headers=headers, json=payload)
        resp.raise_for_status()
        docs = resp.json().get("value") or []
        files = []
        for doc in docs:
            encoded_path = doc.get("metadata_storage_path")
            decoded_url = None
            if encoded_path:
                try:
                    decoded_url = base64.b64decode(encoded_path).decode("utf-8")
                except Exception:
                    decoded_url = encoded_path
                parts = decoded_url.split('/')
                try:
                    container_index = parts.index(blob_container)
                    blob_name = '/'.join(parts[container_index + 1:])
                    blob_name = unquote(blob_name)
                except ValueError:
                    blob_name = None
                file_name = parts[-1] if parts else blob_name
            else:
                blob_name = None
                file_name = "Unknown"
            # Mark as matched if the search query is in the file name (case-insensitive)
            matched = search_query == "*" or (file_name and search_query.lower() in file_name.lower())
            files.append({
                "name": file_name,
                "id": blob_name,
                "matched": matched
            })
        return jsonify({"files": files})
    except Exception as e:
        print("DEBUG ERROR:", e)
        return jsonify({"error": str(e)}), 500

@app.route("/search/download", methods=["POST"])
def download_rfp_file():
    """
    Download a selected RFP file from Azure Blob Storage and analyze directly (no local save).
    """
    data = request.get_json()
    blob_name = data.get("id")
    if not blob_name:
        return jsonify({"error": "No file ID provided"}), 400

    blob_conn_str = os.getenv("AZURE_BLOB_CONNECTION_STRING")
    blob_container = os.getenv("AZURE_BLOB_CONTAINER")

    try:
        blob_service_client = BlobServiceClient.from_connection_string(blob_conn_str)
        blob_client = blob_service_client.get_blob_client(container=blob_container, blob=blob_name)
        blob_data = blob_client.download_blob().readall()

        ext = os.path.splitext(blob_name)[1].lower()
        if ext == ".pdf":
            # Analyze directly from bytes
            result = analyze_layout_from_bytes(blob_data)
            extracted_text = extract_text(result)
            summary = summarize_text(extracted_text) if extracted_text else "No text extracted from document."
            layout = layout_to_text(result)
        elif ext == ".docx":
            # Save DOCX to temp file
            temp_docx = os.path.join(app.config["UPLOAD_FOLDER"], "temp_blob.docx")
            with open(temp_docx, "wb") as f:
                f.write(blob_data)
            # Convert DOCX to PDF
            pdf_path = convert_word_to_pdf(temp_docx)
            # Analyze PDF
            with open(pdf_path, "rb") as f:
                pdf_bytes = f.read()
            result = analyze_layout_from_bytes(pdf_bytes)
            extracted_text = extract_text(result)
            summary = summarize_text(extracted_text) if extracted_text else "No text extracted from document."
            layout = layout_to_text(result)
            # Clean up temp files
            os.remove(temp_docx)
            os.remove(pdf_path)
        else:
            return jsonify({"error": "Unsupported file type. Please select a PDF or DOCX."}), 400

        session['document_text'] = extracted_text
        return jsonify({
            "document_text": extracted_text,
            "summary": summary,
            "layout": layout
        })
    except Exception as e:
        return jsonify({"error": str(e)}), 500

def analyze_layout_from_bytes(pdf_bytes):
    document_intelligence_client = DocumentIntelligenceClient(
        endpoint=AZURE_DOC_INTELLIGENCE_ENDPOINT,
        credential=AzureKeyCredential(AZURE_DOC_INTELLIGENCE_KEY)
    )
    # Analyze directly from bytes
    poller = document_intelligence_client.begin_analyze_document(
        "prebuilt-layout",
        body=pdf_bytes,
        content_type="application/pdf"
    )
    result = poller.result()
    return result

def convert_word_to_pdf(input_path, output_path=None):
    """
    Convert Word document(s) to PDF.
    """
    import os
    if not os.path.exists(input_path):
        raise FileNotFoundError(f"The path '{input_path}' does not exist.")
    try:
        pythoncom.CoInitialize()
        convert(input_path, output_path)
        # If output_path is None, PDF will be in the same directory as input_path
        if output_path:
            pdf_path = output_path
        else:
            pdf_path = os.path.splitext(input_path)[0] + ".pdf"
        return pdf_path
    except Exception as e:
        raise RuntimeError(f"Conversion failed: {e}")

def detect_missing_fields(extracted_text, required_fields=None):
    """
    Checks for missing required fields in the extracted document text.
    Returns a list of missing fields.
    """
    if required_fields is None:
        required_fields = [
            "Signature", "Date", "Name", "Company", "Title"
            # Add more as needed
        ]
    missing = []
    lower_text = extracted_text.lower()
    for field in required_fields:
        if field.lower() not in lower_text:
            missing.append(field)
    return missing

@app.route("/search/query", methods=["POST"])
def semantic_search():
    data = request.get_json()
    user_query = data.get("query")
    if not user_query:
        return jsonify({"error": "No query provided"}), 400

    search_endpoint = os.getenv("AZURE_SEARCH_ENDPOINT")
    search_key = os.getenv("AZURE_SEARCH_KEY")
    search_index = os.getenv("AZURE_SEARCH_INDEX")  # Use your original index
    url = f"{search_endpoint}/indexes/{search_index}/docs/search?api-version=2023-07-01-Preview"
    headers = {
        "Content-Type": "application/json",
        "api-key": search_key
    }
    payload = {
        "search": user_query,
        "queryType": "semantic",
        "queryLanguage": "en-us",
        "top": 3,
        "select": "content,metadata_storage_name,metadata_storage_path",
        "highlight": "content"
    }

    try:
        resp = requests.post(url, headers=headers, json=payload)
        resp.raise_for_status()
        docs = resp.json().get("value", [])
        passages = []
        for doc in docs:
            passages.append({
                "content": doc.get("content", ""),
                "file": doc.get("metadata_storage_name", ""),
                "path": doc.get("metadata_storage_path", ""),
                "highlights": doc.get("@search.highlights", {}).get("content", [])
            })
        return jsonify({"passages": passages})
    except Exception as e:
        return jsonify({"error": str(e)}), 500

if __name__ == "__main__":
    app.run(port=5000, debug=True)

